import re
from docx import Document
from docx.table import Table, _Cell
from typing import Dict, List, Tuple, Any
import json
from datetime import datetime

class SmartFormDetector:
    """
    A smart form detector that can analyze any Word document with tables
    and fill them intelligently based on detected patterns
    """
    
    def __init__(self, template_path: str):
        """Initialize with a Word document path"""
        self.doc = Document(template_path)
        self.tables_info = []
        self.field_patterns = {
            # Personal Information Patterns (Chinese)
            'name': [r'姓\s*名', r'申请人', r'负责人', r'姓名：', r'名字'],
            'gender': [r'性\s*别', r'性别：'],
            'birth_date': [r'出生年月', r'出生日期', r'生日', r'出生时间'],
            'phone': [r'电\s*话', r'联系电话', r'手机', r'联系方式', r'电话号码'],
            'email': [r'邮\s*箱', r'电子邮件', r'E-?mail', r'电邮'],
            'ethnicity': [r'民\s*族', r'民族：'],
            'nationality': [r'国\s*籍', r'国籍：'],
            'id_number': [r'身份证', r'证件号', r'身份证号'],
            'address': [r'地\s*址', r'住址', r'通讯地址', r'联系地址'],
            
            # Professional Information
            'title': [r'职\s*称', r'职务', r'专业技术职务', r'技术职称', r'职位'],
            'department': [r'部\s*门', r'院系', r'单位', r'所在部门', r'工作单位'],
            'degree': [r'学\s*位', r'学历', r'最高学位', r'最终学位'],
            'major': [r'专\s*业', r'研究方向', r'研究领域', r'专业方向'],
            
            # Project/Research Related
            'project_name': [r'项目名称', r'课题名称', r'研究名称', r'项目题目'],
            'project_number': [r'项目编号', r'项目号', r'课题编号', r'编号'],
            'funding': [r'经\s*费', r'资助金额', r'项目经费', r'资金', r'金额'],
            'period': [r'期\s*限', r'周期', r'起止时间', r'时间段', r'年限'],
            'date': [r'日\s*期', r'时间', r'年月日', r'申请日期'],
            
            # Publication Related
            'paper_title': [r'论文题目', r'论文名称', r'文章标题', r'论文标题', r'题目'],
            'journal': [r'期\s*刊', r'杂志', r'发表期刊', r'刊物', r'会议'],
            'author': [r'作\s*者', r'著者', r'作者姓名', r'第一作者'],
            
            # Award Related
            'award_name': [r'奖项名称', r'获奖名称', r'奖励名称', r'成果名称'],
            'award_level': [r'奖励等级', r'获奖等级', r'奖项级别', r'等级'],
            'award_date': [r'获奖时间', r'获奖日期', r'颁奖时间'],
            
            # Other Common Fields
            'description': [r'描\s*述', r'简介', r'说明', r'内容', r'详情', r'成果简介'],
            'notes': [r'备\s*注', r'说明', r'其他', r'附注', r'注释'],
            'signature': [r'签\s*名', r'签字', r'申请人签名', r'负责人签名'],
            'year': [r'年\s*度', r'年份', r'学年', r'年'],
        }
        
        # Analyze all tables in the document
        self.analyze_document()
    
    def analyze_document(self):
        """Analyze all tables in the document"""
        print(f"Found {len(self.doc.tables)} tables in the document")
        
        for idx, table in enumerate(self.doc.tables):
            table_info = {
                'index': idx,
                'rows': len(table.rows),
                'cols': len(table.columns),
                'headers': [],
                'field_mapping': {},
                'structure': self.analyze_table_structure(table)
            }
            
            # Try to identify headers and fields
            self.identify_table_fields(table, table_info)
            self.tables_info.append(table_info)
            
            print(f"\nTable {idx + 1}:")
            print(f"  Dimensions: {table_info['rows']} rows × {table_info['cols']} columns")
            print(f"  Detected fields: {list(table_info['field_mapping'].keys())}")
    
    def analyze_table_structure(self, table: Table) -> str:
        """Determine the structure type of the table"""
        if len(table.rows) == 0:
            return "empty"
        
        # Check if it's a vertical form (label-value pairs in rows)
        first_col_texts = [row.cells[0].text.strip() for row in table.rows if len(row.cells) > 0]
        if self.contains_multiple_labels(first_col_texts):
            return "vertical"
        
        # Check if it's a horizontal form (headers in first row)
        if len(table.rows) > 0:
            first_row_texts = [cell.text.strip() for cell in table.rows[0].cells]
            if self.contains_multiple_labels(first_row_texts):
                return "horizontal"
        
        # Check if it's a mixed/complex form
        return "mixed"
    
    def contains_multiple_labels(self, texts: List[str]) -> bool:
        """Check if texts contain multiple field labels"""
        label_count = 0
        for text in texts:
            for field_type, patterns in self.field_patterns.items():
                for pattern in patterns:
                    if re.search(pattern, text, re.IGNORECASE):
                        label_count += 1
                        break
        return label_count >= 2
    
    def identify_table_fields(self, table: Table, table_info: Dict):
        """Identify fields in a table based on patterns"""
        structure = table_info['structure']
        
        if structure == "vertical":
            self.identify_vertical_fields(table, table_info)
        elif structure == "horizontal":
            self.identify_horizontal_fields(table, table_info)
        else:
            self.identify_mixed_fields(table, table_info)
    
    def identify_vertical_fields(self, table: Table, table_info: Dict):
        """Identify fields in vertical tables (label in one column, value in next)"""
        for row_idx, row in enumerate(table.rows):
            for col_idx in range(len(row.cells) - 1):
                label_text = row.cells[col_idx].text.strip()
                
                # Check if this cell contains a field label
                for field_type, patterns in self.field_patterns.items():
                    for pattern in patterns:
                        if re.search(pattern, label_text, re.IGNORECASE):
                            # The next cell is likely the value field
                            table_info['field_mapping'][field_type] = {
                                'label_cell': (row_idx, col_idx),
                                'value_cell': (row_idx, col_idx + 1),
                                'label_text': label_text,
                                'type': 'vertical'
                            }
                            break
    
    def identify_horizontal_fields(self, table: Table, table_info: Dict):
        """Identify fields in horizontal tables (headers in first row)"""
        if len(table.rows) < 2:
            return
        
        # Analyze headers
        headers = []
        for col_idx, cell in enumerate(table.rows[0].cells):
            header_text = cell.text.strip()
            headers.append(header_text)
            
            # Check if this header matches any field pattern
            for field_type, patterns in self.field_patterns.items():
                for pattern in patterns:
                    if re.search(pattern, header_text, re.IGNORECASE):
                        table_info['field_mapping'][f"{field_type}_{col_idx}"] = {
                            'label_cell': (0, col_idx),
                            'value_cells': [(row_idx, col_idx) for row_idx in range(1, len(table.rows))],
                            'label_text': header_text,
                            'type': 'horizontal'
                        }
                        break
        
        table_info['headers'] = headers
    
    def identify_mixed_fields(self, table: Table, table_info: Dict):
        """Identify fields in mixed/complex tables"""
        # Search all cells for patterns
        for row_idx, row in enumerate(table.rows):
            for col_idx, cell in enumerate(row.cells):
                cell_text = cell.text.strip()
                
                # Skip empty cells
                if not cell_text:
                    continue
                
                # Check if this cell contains a field label
                for field_type, patterns in self.field_patterns.items():
                    for pattern in patterns:
                        if re.search(pattern, cell_text, re.IGNORECASE):
                            # Determine where the value might be
                            value_location = self.find_value_location(table, row_idx, col_idx)
                            if value_location:
                                key = f"{field_type}_{row_idx}_{col_idx}"
                                table_info['field_mapping'][key] = {
                                    'label_cell': (row_idx, col_idx),
                                    'value_cell': value_location,
                                    'label_text': cell_text,
                                    'type': 'mixed'
                                }
                            break
    
    def find_value_location(self, table: Table, label_row: int, label_col: int) -> Tuple[int, int]:
        """Find the most likely location for a value given a label location"""
        # Check right cell first
        if label_col + 1 < len(table.rows[label_row].cells):
            right_cell = table.rows[label_row].cells[label_col + 1].text.strip()
            if not self.is_likely_label(right_cell):
                return (label_row, label_col + 1)
        
        # Check cell below
        if label_row + 1 < len(table.rows) and label_col < len(table.rows[label_row + 1].cells):
            below_cell = table.rows[label_row + 1].cells[label_col].text.strip()
            if not self.is_likely_label(below_cell):
                return (label_row + 1, label_col)
        
        # Check same cell (label and value might be in same cell with separator)
        return (label_row, label_col)
    
    def is_likely_label(self, text: str) -> bool:
        """Check if text is likely a label rather than a value"""
        # Check if it matches any label pattern
        for patterns in self.field_patterns.values():
            for pattern in patterns:
                if re.search(pattern, text, re.IGNORECASE):
                    return True
        
        # Check for common label indicators
        label_indicators = ['：', ':', '(', '（', '/', '、']
        return any(indicator in text for indicator in label_indicators)
    
    def get_cell_value(self, table: Table, row: int, col: int) -> str:
        """Safely get cell value"""
        try:
            return table.rows[row].cells[col].text.strip()
        except:
            return ""
    
    def set_cell_value(self, table: Table, row: int, col: int, value: str):
        """Safely set cell value"""
        try:
            cell = table.rows[row].cells[col]
            # Clear existing content
            for paragraph in cell.paragraphs:
                paragraph.clear()
            # Add new content
            paragraph = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
            paragraph.text = value
        except Exception as e:
            print(f"Error setting cell ({row}, {col}): {e}")
    
    def fill_form(self, data: Dict[str, Any]):
        """
        Fill the form with provided data
        
        Args:
            data: Dictionary with field types as keys and values to fill
                  e.g., {'name': 'John Doe', 'phone': '123-456-7890'}
        """
        filled_count = 0
        
        for table_idx, table_info in enumerate(self.tables_info):
            table = self.doc.tables[table_idx]
            
            print(f"\nFilling Table {table_idx + 1}:")
            
            for field_key, field_info in table_info['field_mapping'].items():
                # Extract base field type (remove indices)
                base_field_type = field_key.split('_')[0]
                
                if base_field_type in data:
                    value = str(data[base_field_type])
                    
                    if field_info['type'] == 'vertical':
                        # Fill single cell
                        row, col = field_info['value_cell']
                        self.set_cell_value(table, row, col, value)
                        print(f"  Filled {field_info['label_text']}: {value}")
                        filled_count += 1
                    
                    elif field_info['type'] == 'horizontal':
                        # For horizontal tables, we might need to handle multiple rows
                        if 'value_cells' in field_info and field_info['value_cells']:
                            # Fill first available row
                            row, col = field_info['value_cells'][0]
                            self.set_cell_value(table, row, col, value)
                            print(f"  Filled {field_info['label_text']}: {value}")
                            filled_count += 1
                    
                    elif field_info['type'] == 'mixed':
                        row, col = field_info['value_cell']
                        
                        # Check if value is in the same cell as label
                        if (row, col) == field_info['label_cell']:
                            # Replace or append value after label
                            current_text = self.get_cell_value(table, row, col)
                            new_text = f"{field_info['label_text']}：{value}"
                            self.set_cell_value(table, row, col, new_text)
                        else:
                            self.set_cell_value(table, row, col, value)
                        
                        print(f"  Filled {field_info['label_text']}: {value}")
                        filled_count += 1
        
        print(f"\nTotal fields filled: {filled_count}")
        return filled_count
    
    def save(self, output_path: str):
        """Save the filled document"""
        self.doc.save(output_path)
        print(f"Document saved to: {output_path}")
    
    def export_field_mapping(self, output_file: str = "field_mapping.json"):
        """Export the detected field mapping for review"""
        export_data = []
        
        for table_idx, table_info in enumerate(self.tables_info):
            table_data = {
                'table_index': table_idx,
                'dimensions': f"{table_info['rows']}x{table_info['cols']}",
                'structure': table_info['structure'],
                'fields': {}
            }
            
            for field_key, field_info in table_info['field_mapping'].items():
                table_data['fields'][field_key] = {
                    'label': field_info['label_text'],
                    'type': field_info['type'],
                    'location': str(field_info.get('value_cell', field_info.get('value_cells', 'N/A')))
                }
            
            export_data.append(table_data)
        
        with open(output_file, 'w', encoding='utf-8') as f:
            json.dump(export_data, f, ensure_ascii=False, indent=2)
        
        print(f"Field mapping exported to: {output_file}")


# Example usage and testing
if __name__ == "__main__":
    import sys
    
    # Get document path from command line or use default
    if len(sys.argv) > 1:
        doc_path = sys.argv[1]
    else:
        doc_path = input("Enter the path to your Word document: ").strip()
    
    try:
        # Create detector instance
        detector = SmartFormDetector(doc_path)
        
        # Export field mapping for review
        detector.export_field_mapping()
        
        # Example data to fill
        sample_data = {
            'name': '张三',
            'gender': '男',
            'birth_date': '1990年1月1日',
            'phone': '13800138000',
            'email': 'zhangsan@example.com',
            'ethnicity': '汉',
            'title': '教授',
            'department': '计算机科学系',
            'degree': '博士',
            'major': '人工智能',
            'project_name': '基于深度学习的自然语言处理研究',
            'funding': '100万元',
            'period': '2023-2026',
            'description': '本项目旨在研究最新的深度学习技术在自然语言处理中的应用...',
        }
        
        print("\n" + "="*50)
        print("Sample data to fill:")
        for key, value in sample_data.items():
            print(f"  {key}: {value}")
        
        print("\n" + "="*50)
        choice = input("\nDo you want to fill the form with sample data? (y/n): ")
        
        if choice.lower() == 'y':
            # Fill the form
            detector.fill_form(sample_data)
            
            # Save the result
            output_path = doc_path.replace('.docx', '_filled.docx')
            detector.save(output_path)
        else:
            print("\nYou can create your own data dictionary and call:")
            print("  detector.fill_form(your_data)")
            print("  detector.save('output.docx')")
        
    except Exception as e:
        print(f"Error: {e}")
        import traceback
        traceback.print_exc()


#how to use
# Create detector for any Word document
detector = SmartFormDetector("any_form.docx")

# Prepare your data
data = {
    'name': 'John Doe',
    'phone': '123-456-7890',
    'email': 'john@example.com',
    'title': 'Professor',
    # Add any fields you want to fill
}

# Fill the form
detector.fill_form(data)

# Save the result
detector.save("filled_form.docx")

#yunxing
#python smart_form_detector.py your_document.docx


# Add your own patterns for specific fields
detector.field_patterns['custom_field'] = [r'特殊字段', r'Special Field']

# For tables with multiple rows (like publications)
data = {
    'paper_title': 'My Research Paper',  # Will fill first available row
}

# See what fields were detected
detector.export_field_mapping("my_fields.json")


# | Name    | [John]  |
# | Phone   | [12345] |

# | Name | Phone | Email |
# | [  ] | [   ] | [   ] |

# | Personal Info          |
# | Name: [ ] Gender: [ ]  |
# | Phone: [ ] Email: [ ]  |