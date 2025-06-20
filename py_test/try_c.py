import pandas as pd
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

class ResearchAwardFormFiller:
    def __init__(self, template_path):
        """
        Initialize the form filler with a template document
        
        Args:
            template_path: Path to the template Word document
        """
        self.doc = Document(template_path)
        self.form_data = {
            # Table 1: Basic Information (基本信息)
            "basic_info": {
                "name": "",  # 姓名
                "gender": "",  # 性别
                "ethnicity": "",  # 民族
                "birth_date": "",  # 出生年月
                "professional_title": "",  # 专业技术职务
                "administrative_position": "",  # 行政职务
                "department_head": "",  # 系主任
                "final_degree": "",  # 最终学位及授予国家或地区及学校
                "research_direction": "",  # 研究方向
                "contact_phone": "",  # 联系电话
                "work_unit": "",  # 所在工作单位
                "academic_positions": [],  # 主要学术任职 (list of strings)
            },
            
            # Table 2: Main Innovations (主要创新成果)
            "innovations": "",  # Long text describing innovations
            
            # Table 3: National Projects (国家级重点项目)
            "projects": [
                # Each project is a dict with keys: "number", "name", "funding", "period", "source"
            ],
            
            # Table 4: Publications (重要著作和论文)
            "publications": [
                # Each publication is a dict with keys: "title", "date", "journal", "volume", "level", "author_rank"
            ],
            
            # Table 5: Think Tank Results (智库成果)
            "think_tank": [
                # Each result is a dict with keys: "name", "year", "adopting_unit", "author_rank", "notes"
            ],
            
            # Table 6: Patents (授权发明专利)
            "patents": [
                # Each patent is a dict with keys: "name", "patent_number", "year", "country", "author_rank", "economic_benefit"
            ],
            
            # Table 7: Awards (获奖成果)
            "awards": [
                # Each award is a dict with keys: "project_name", "award_type", "granting_unit", "award_date"
            ]
        }
    
    def set_basic_info(self, **kwargs):
        """Set basic information fields"""
        for key, value in kwargs.items():
            if key in self.form_data["basic_info"]:
                self.form_data["basic_info"][key] = value
    
    def set_innovations(self, text):
        """Set innovations description"""
        self.form_data["innovations"] = text
    
    def add_project(self, number, name, funding, period, source):
        """Add a project to the projects list"""
        self.form_data["projects"].append({
            "number": number,
            "name": name,
            "funding": funding,
            "period": period,
            "source": source
        })
    
    def add_publication(self, title, date, journal, volume="", level="", author_rank=""):
        """Add a publication to the publications list"""
        self.form_data["publications"].append({
            "title": title,
            "date": date,
            "journal": journal,
            "volume": volume,
            "level": level,
            "author_rank": author_rank
        })
    
    def add_think_tank_result(self, name, year, adopting_unit, author_rank="", notes=""):
        """Add a think tank result"""
        self.form_data["think_tank"].append({
            "name": name,
            "year": year,
            "adopting_unit": adopting_unit,
            "author_rank": author_rank,
            "notes": notes
        })
    
    def add_patent(self, name, patent_number, year, country, author_rank, economic_benefit=""):
        """Add a patent"""
        self.form_data["patents"].append({
            "name": name,
            "patent_number": patent_number,
            "year": year,
            "country": country,
            "author_rank": author_rank,
            "economic_benefit": economic_benefit
        })
    
    def add_award(self, project_name, award_type, granting_unit, award_date):
        """Add an award"""
        self.form_data["awards"].append({
            "project_name": project_name,
            "award_type": award_type,
            "granting_unit": granting_unit,
            "award_date": award_date
        })
    
    def fill_table_cell(self, table, row_idx, col_idx, text):
        """Helper function to fill a specific cell in a table"""
        try:
            cell = table.rows[row_idx].cells[col_idx]
            # Clear existing content
            for paragraph in cell.paragraphs:
                paragraph.clear()
            # Add new content
            paragraph = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
            paragraph.text = str(text)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        except Exception as e:
            print(f"Error filling cell at row {row_idx}, col {col_idx}: {e}")
    
    def fill_form(self):
        """Fill the form with the provided data"""
        tables = self.doc.tables
        
        # Fill Table 1: Basic Information
        if len(tables) > 0:
            table1 = tables[0]
            info = self.form_data["basic_info"]
            
            # Map fields to their positions in the table
            # Note: These positions are estimates based on the document structure
            # You may need to adjust them based on the actual table layout
            
            # Name (姓名)
            self.fill_table_cell(table1, 1, 2, info["name"])
            # Gender (性别)
            self.fill_table_cell(table1, 1, 4, info["gender"])
            # Ethnicity (民族)
            self.fill_table_cell(table1, 1, 6, info["ethnicity"])
            # Birth date (出生年月)
            self.fill_table_cell(table1, 1, 8, info["birth_date"])
            
            # Professional title (专业技术职务)
            self.fill_table_cell(table1, 2, 2, info["professional_title"])
            # Final degree (最终学位)
            self.fill_table_cell(table1, 2, 8, info["final_degree"])
            
            # Research direction (研究方向)
            self.fill_table_cell(table1, 3, 2, info["research_direction"])
            # Contact phone (联系电话)
            self.fill_table_cell(table1, 3, 8, info["contact_phone"])
            
            # Work unit (所在工作单位)
            self.fill_table_cell(table1, 4, 4, info["work_unit"])
            
            # Academic positions (主要学术任职)
            if info["academic_positions"]:
                positions_text = "\n".join(info["academic_positions"])
                self.fill_table_cell(table1, 5, 1, positions_text)
        
        # Fill Table 2: Innovations
        if len(tables) > 1:
            table2 = tables[1]
            self.fill_table_cell(table2, 1, 0, self.form_data["innovations"])
        
        # Fill Table 3: Projects
        if len(tables) > 2 and self.form_data["projects"]:
            table3 = tables[2]
            for i, project in enumerate(self.form_data["projects"], start=1):
                if i < len(table3.rows):
                    self.fill_table_cell(table3, i, 0, project["number"])
                    self.fill_table_cell(table3, i, 1, project["name"])
                    self.fill_table_cell(table3, i, 2, project["funding"])
                    self.fill_table_cell(table3, i, 3, project["period"])
                    self.fill_table_cell(table3, i, 4, project["source"])
        
        # Fill Table 4: Publications
        if len(tables) > 3 and self.form_data["publications"]:
            table4 = tables[3]
            for i, pub in enumerate(self.form_data["publications"], start=1):
                if i < len(table4.rows):
                    self.fill_table_cell(table4, i, 0, pub["title"])
                    self.fill_table_cell(table4, i, 1, pub["date"])
                    self.fill_table_cell(table4, i, 2, pub["journal"])
                    self.fill_table_cell(table4, i, 3, pub["volume"])
                    self.fill_table_cell(table4, i, 4, pub["level"])
                    self.fill_table_cell(table4, i, 5, pub["author_rank"])
        
        # Fill Table 5: Think Tank Results
        if len(tables) > 4 and self.form_data["think_tank"]:
            table5 = tables[4]
            for i, result in enumerate(self.form_data["think_tank"], start=1):
                if i < len(table5.rows):
                    self.fill_table_cell(table5, i, 0, result["name"])
                    self.fill_table_cell(table5, i, 1, result["year"])
                    self.fill_table_cell(table5, i, 2, result["adopting_unit"])
                    self.fill_table_cell(table5, i, 3, result["author_rank"])
                    self.fill_table_cell(table5, i, 4, result["notes"])
        
        # Fill Table 6: Patents
        if len(tables) > 5 and self.form_data["patents"]:
            table6 = tables[5]
            for i, patent in enumerate(self.form_data["patents"], start=1):
                if i < len(table6.rows):
                    self.fill_table_cell(table6, i, 0, patent["name"])
                    self.fill_table_cell(table6, i, 1, patent["patent_number"])
                    self.fill_table_cell(table6, i, 2, patent["year"])
                    self.fill_table_cell(table6, i, 3, patent["country"])
                    self.fill_table_cell(table6, i, 4, patent["author_rank"])
                    self.fill_table_cell(table6, i, 5, patent["economic_benefit"])
        
        # Fill Table 7: Awards
        if len(tables) > 6 and self.form_data["awards"]:
            table7 = tables[6]
            for i, award in enumerate(self.form_data["awards"], start=1):
                if i < len(table7.rows):
                    self.fill_table_cell(table7, i, 0, award["project_name"])
                    self.fill_table_cell(table7, i, 1, award["award_type"])
                    self.fill_table_cell(table7, i, 2, award["granting_unit"])
                    self.fill_table_cell(table7, i, 3, award["award_date"])
    
    def save(self, output_path):
        """Save the filled document"""
        self.fill_form()
        self.doc.save(output_path)
        print(f"Document saved to: {output_path}")


# Example usage:
if __name__ == "__main__":
    # Create form filler instance
    filler = ResearchAwardFormFiller("/home/lsyedith/py_test/empty_list.docx")  # Replace with your template path
    
    # Fill basic information
    filler.set_basic_info(
        name="张三",
        gender="男",
        ethnicity="汉",
        birth_date="1985年1月",
        professional_title="教授",
        administrative_position="",
        department_head="否",
        final_degree="博士，北京大学",
        research_direction="人工智能与机器学习",
        contact_phone="13800138000",
        work_unit="中国人民大学信息学院",
        academic_positions=[
            "- IEEE会员",
            "- ACM会员",
            "- 中国计算机学会高级会员",
            "- 某期刊编委"
        ]
    )
    
    # Set innovations
    filler.set_innovations("""
    在过去三年中，本人在人工智能领域取得了以下创新成果：
    1. 提出了新的深度学习算法...
    2. 开发了创新的自然语言处理模型...
    3. 在计算机视觉方面取得突破...
    """)
    
    # Add projects
    filler.add_project("1", "国家自然科学基金项目", "100", "2022-2025", "国家自然科学基金委")
    filler.add_project("2", "科技部重点研发计划", "500", "2023-2026", "科技部")
    
    # Add publications
    filler.add_publication(
        "Deep Learning for NLP: A Survey",
        "2023年6月",
        "Nature Machine Intelligence",
        "5(6)",
        "学校A+ CCF A",
        "第一作者"
    )
    
    # Add awards
    filler.add_award("优秀科研成果", "省部级一等奖", "教育部", "2023")
    
    # Save the filled form
    filler.save("/home/lsyedith/py_test/filled_form.docx")




# Initialize with your template document
filler = ResearchAwardFormFiller("/home/lsyedith/py_test/empty_list.docx")

# Fill in the information
filler.set_basic_info(
    name="Your Name",
    gender="Male/Female",
    # ... other fields
)

# Add entries to tables
filler.add_publication("Paper Title", "2023年6月", "Journal Name", "Vol(Issue)", "CCF A", "First Author")

# Save the filled form
filler.save("/home/lsyedith/py_test/output.docx")